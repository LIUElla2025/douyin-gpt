"""Vercel Serverless API — 稿了个 AI（视频文稿提取 & 知识分身）

Flask 应用，提供以下 API：
- POST /api/resolve-url    解析抖音链接 → sec_uid
- POST /api/fetch-videos   获取博主视频列表
- POST /api/transcribe     转录单个视频音频
- POST /api/generate-doc   生成 Word 文档
- POST /api/upload-corpus  上传文档创建语料库
- POST /api/chat           博主 GPT 对话 / 语料库对话
"""

import base64
import hashlib
import json
import os
import random
import re
import string
import tempfile
import time
import urllib.parse
import urllib.request
from datetime import datetime
from io import BytesIO

from flask import Flask, Response, jsonify, request, send_file

app = Flask(__name__)

# ─── 从 Vercel 环境变量读取默认配置 ───
_ENV_OPENAI_KEY = os.environ.get("OPENAI_API_KEY", "").strip()
_ENV_COOKIE = os.environ.get("DOUYIN_COOKIE", "").strip()
_ENV_APIFY_TOKEN = os.environ.get("APIFY_API_TOKEN", "").strip()
_ENV_PROXY = os.environ.get("PROXY", "").strip()


def _clean_cookie(cookie: str) -> str:
    """清理 Cookie 中的非法字符（换行、回车等）"""
    return re.sub(r'[\r\n\t]+', '', cookie).strip()


def _get_config(data: dict) -> dict:
    """合并请求参数和环境变量，环境变量作为默认值"""
    return {
        "openai_api_key": data.get("openai_api_key", "").strip() or _ENV_OPENAI_KEY,
        "cookie": _clean_cookie(data.get("cookie", "") or _ENV_COOKIE),
        "apify_token": data.get("apify_token", "").strip() or _ENV_APIFY_TOKEN,
        "proxy": data.get("proxy", "").strip() or _ENV_PROXY,
    }


# ─── CORS ───


@app.after_request
def add_cors(response):
    response.headers["Access-Control-Allow-Origin"] = "*"
    response.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
    response.headers["Access-Control-Allow-Headers"] = "Content-Type"
    return response


@app.route("/api/<path:path>", methods=["OPTIONS"])
def handle_options(path):
    return "", 204


# ─── 健康检查 ───


@app.route("/api/health", methods=["GET"])
def health():
    return jsonify({"status": "ok", "timestamp": datetime.utcnow().isoformat()})


@app.route("/api/config-status", methods=["GET"])
def config_status():
    """返回服务端已配置哪些 Key（不暴露值，只返回 bool）"""
    return jsonify({
        "has_openai_key": bool(_ENV_OPENAI_KEY),
        "has_cookie": bool(_ENV_COOKIE),
        "has_apify_token": bool(_ENV_APIFY_TOKEN),
        "has_proxy": bool(_ENV_PROXY),
    })


# ─── 解析抖音链接 ───


@app.route("/api/resolve-url", methods=["POST"])
def resolve_url():
    data = request.json or {}
    user_input = data.get("input", "").strip()
    if not user_input:
        return jsonify({"error": "请输入抖音链接或抖音号"}), 400

    try:
        profile_url = _resolve_douyin_input(user_input)
        sec_uid = _extract_sec_uid(profile_url)
        return jsonify({"profile_url": profile_url, "sec_uid": sec_uid})
    except Exception as e:
        return jsonify({"error": f"解析失败: {e}"}), 400


# ─── 获取视频列表（流式 SSE） ───


@app.route("/api/fetch-videos", methods=["POST"])
def fetch_videos():
    data = request.json or {}
    cfg = _get_config(data)
    sec_uid = data.get("sec_uid", "").strip()
    cookie = cfg["cookie"]
    keyword = data.get("keyword", "").strip()
    max_videos = data.get("max_videos", 0)
    # 目标 ID 模式：只找这些视频，找齐就停
    target_ids = set(data.get("target_ids", []))
    # 续传参数：从上次超时的位置继续扫描
    start_cursor = data.get("start_cursor", 0)
    start_page = data.get("start_page", 0)
    prev_scanned = data.get("prev_scanned", 0)
    prev_matched = data.get("prev_matched", 0)

    if not sec_uid:
        return jsonify({"error": "缺少 sec_uid"}), 400
    if not cookie:
        return jsonify({"error": "请配置抖音 Cookie"}), 400

    def sse_generate():
        """SSE 流式生成：逐页推送视频列表，前端实时显示"""

        def send_event(event_type, data_dict):
            return f"event: {event_type}\ndata: {json.dumps(data_dict, ensure_ascii=False)}\n\n"

        try:
            # 第一步：获取用户信息
            yield send_event("status", {"msg": "正在获取博主信息..."})
            profile_data = _fetch_user_profile_direct(sec_uid, cookie)
            user = profile_data.get("user", {})
            creator_name = user.get("nickname", "")
            total_videos = user.get("aweme_count", 0)

            yield send_event("profile", {
                "creator_name": creator_name,
                "total_videos": total_videos,
            })

            # 第二步：逐页获取视频（支持从 start_cursor 续传）
            keywords = keyword.split() if keyword else None
            all_videos = []
            total_scanned = prev_scanned
            max_cursor = start_cursor
            max_count = len(target_ids) if target_ids else (max_videos if max_videos > 0 else 9999)
            page = start_page
            total_pages = max(1, -(-total_videos // 35))  # ceil division
            fetch_start = time.time()

            if start_cursor:
                yield send_event("status", {
                    "msg": f"从第 {page + 1} 页续传扫描...",
                })

            stop_reason = "已扫描全部视频"
            while len(all_videos) < max_count:
                page += 1
                yield send_event("status", {
                    "msg": f"正在扫描第 {page}/{total_pages} 页（已扫描 {total_scanned}，匹配 {prev_matched + len(all_videos)}）",
                    "fetched": prev_matched + len(all_videos),
                    "total": total_videos,
                    "page": page,
                    "total_pages": total_pages,
                })

                params = _build_base_params(sec_uid, max_cursor, count=35)
                page_data = _douyin_api_request(
                    "https://www.douyin.com/aweme/v1/web/aweme/post/", params, cookie
                )

                aweme_list = page_data.get("aweme_list", [])
                if not aweme_list:
                    stop_reason = f"第{page}页返回空数据"
                    break

                page_videos = []
                total_scanned += len(aweme_list)
                for item in aweme_list:
                    vid = item.get("aweme_id", "")
                    desc = item.get("desc", "无标题")
                    title = re.sub(r"#\S+", "", desc).strip() or "无标题"
                    duration = item.get("video", {}).get("duration", 0)
                    if isinstance(duration, (int, float)) and duration > 10000:
                        duration = duration // 1000
                    # 视频下载URL：优先用 download_addr（完整MP4，含音频+视频），
                    # play_addr 是 DASH fMP4 格式，通常只含视频轨道无音频
                    video_obj = item.get("video", {})
                    video_download_url = ""
                    # 方式1: download_addr（完整MP4，含音频轨道）
                    download_addr = video_obj.get("download_addr", {})
                    da_list = download_addr.get("url_list", [])
                    if da_list:
                        video_download_url = da_list[0]
                    # 方式2: play_addr（DASH fMP4，可能仅含视频轨道）
                    if not video_download_url:
                        play_addr = video_obj.get("play_addr", {})
                        pa_list = play_addr.get("url_list", [])
                        if pa_list:
                            video_download_url = pa_list[0]
                    # 方式3: 最低码率（备选）
                    if not video_download_url:
                        bit_rates = video_obj.get("bit_rate", [])
                        if bit_rates:
                            lowest = min(bit_rates, key=lambda b: b.get("bit_rate", float("inf")))
                            br_addr = lowest.get("play_addr", {})
                            br_list = br_addr.get("url_list", [])
                            video_download_url = br_list[0] if br_list else ""
                    # 背景音乐URL（最后备选）
                    music = item.get("music", {})
                    play_url_list = (music.get("play_url") or {}).get("url_list", [])
                    audio_url = play_url_list[0] if play_url_list else ""
                    author = (item.get("author") or {}).get("nickname", creator_name)
                    create_time = item.get("create_time", "")

                    video = {
                        "id": str(vid),
                        "title": title,
                        "raw_title": desc,
                        "url": f"https://www.douyin.com/video/{vid}",
                        "create_time": create_time,
                        "duration": duration,
                        "author": author,
                        "video_download_url": video_download_url,
                        "audio_url": audio_url,
                        "creator_name": creator_name,
                    }

                    if target_ids:
                        # target_ids 模式：只匹配目标 ID，跳过关键词过滤
                        if video["id"] in target_ids:
                            page_videos.append(video)
                    elif keywords:
                        if _match_keyword(video, keywords):
                            page_videos.append(video)
                    else:
                        page_videos.append(video)

                all_videos.extend(page_videos)

                # 推送本页获取到的视频
                if page_videos:
                    yield send_event("videos", {
                        "videos": page_videos,
                        "fetched": prev_matched + len(all_videos),
                        "total": total_videos,
                    })

                # target_ids 模式：找齐目标视频立即停止
                if target_ids:
                    found_ids = {v["id"] for v in all_videos}
                    remaining = target_ids - found_ids
                    if not remaining:
                        stop_reason = f"已找到全部 {len(target_ids)} 个目标视频"
                        break

                if len(all_videos) >= max_count:
                    stop_reason = f"已达到最大数量限制({max_count})"
                    break

                has_more = page_data.get("has_more", False)
                max_cursor = page_data.get("max_cursor", 0)
                if not has_more:
                    stop_reason = f"API返回has_more=false（已到末页）"
                    break
                if not max_cursor:
                    stop_reason = f"API返回max_cursor=0（无下一页）"
                    break

                time.sleep(0.3)

                # 超时保护：250秒后停止，返回游标供前端续传
                if time.time() - fetch_start > 250:
                    stop_reason = "timeout"
                    yield send_event("status", {"msg": f"本轮超时，已扫描 {total_scanned} 个，准备续传..."})
                    break

            # 完成（包含续传信息）
            done_data = {
                "videos": all_videos,
                "creator_name": creator_name,
                "total": len(all_videos),
                "total_scanned": total_scanned,
                "total_pages_scanned": page,
                "stop_reason": stop_reason,
                "method": "direct",
            }
            # 超时时返回续传游标
            if stop_reason == "timeout":
                done_data["resume_cursor"] = max_cursor
                done_data["resume_page"] = page
            yield send_event("done", done_data)

        except Exception as e:
            yield send_event("error", {"error": str(e)})

    return Response(
        sse_generate(),
        mimetype="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Access-Control-Allow-Origin": "*",
        },
    )


# ─── 转录单个视频 ───


@app.route("/api/transcribe", methods=["POST"])
def transcribe():
    data = request.json or {}
    cfg = _get_config(data)
    audio_url = data.get("audio_url", "").strip()
    video_download_url = data.get("video_download_url", "").strip()
    video_url = data.get("video_url", "").strip()
    video_id = data.get("video_id", "").strip()
    openai_key = cfg["openai_api_key"]
    proxy = cfg["proxy"]
    cookie = cfg["cookie"]

    if not openai_key:
        return jsonify({"error": "缺少 OpenAI API Key"}), 400

    download_url = audio_url or video_download_url or video_url
    if not download_url:
        return jsonify({"error": "缺少音频/视频 URL"}), 400

    tmp_files = []

    try:
        # --- 第1步：下载 ---
        # 有 audio_url 时下载小音频文件；否则下载视频但限制 25MB
        is_audio = bool(audio_url)
        suffix = ".mp3" if is_audio else ".mp4"
        max_dl = 0 if is_audio else 24 * 1024 * 1024  # 视频只下前24MB

        tmp = tempfile.NamedTemporaryFile(suffix=suffix, delete=False)
        tmp_path = tmp.name
        tmp.close()
        tmp_files.append(tmp_path)

        _download_url(download_url, tmp_path, cookie, proxy, max_bytes=max_dl)

        file_size = os.path.getsize(tmp_path)
        if file_size < 1000:
            # 主URL失败，尝试备用
            fallback = video_download_url if download_url == audio_url else audio_url
            if fallback and fallback != download_url:
                is_fb_audio = (fallback == audio_url) if audio_url else False
                new_suffix = ".mp3" if is_fb_audio else ".mp4"
                fb_max = 0 if is_fb_audio else 24 * 1024 * 1024
                new_tmp = tempfile.NamedTemporaryFile(suffix=new_suffix, delete=False)
                new_tmp_path = new_tmp.name
                new_tmp.close()
                tmp_files.append(new_tmp_path)
                _download_url(fallback, new_tmp_path, cookie, proxy, max_bytes=fb_max)
                if os.path.getsize(new_tmp_path) >= 1000:
                    tmp_path = new_tmp_path
                    file_size = os.path.getsize(tmp_path)
            if file_size < 1000:
                return jsonify({"error": f"音频文件太小({file_size}字节)，URL可能已过期"}), 400

        # --- 第2步：准备Whisper输入 ---
        whisper_file = tmp_path

        # 对 mp4，尝试用 ffmpeg 提取纯音频（大幅缩小文件）
        if tmp_path.endswith(".mp4"):
            try:
                audio_path = _extract_audio(tmp_path)
                tmp_files.append(audio_path)
                whisper_file = audio_path
            except Exception:
                # ffmpeg 不可用，直接用截断的 mp4（Whisper 支持 mp4）
                whisper_file = tmp_path

        # 最终大小检查（音频文件不应超过25MB，视频已截断到24MB）
        whisper_size = os.path.getsize(whisper_file)
        if whisper_size > 25 * 1024 * 1024:
            return jsonify({"error": f"音频过大({whisper_size // 1024 // 1024}MB)，Whisper限制25MB"}), 400

        # --- 第3步：调用 Whisper ---
        transcript = _call_whisper(whisper_file, openai_key, proxy)

        # --- GPT 后处理加标点 ---
        raw_text = transcript.get("text", "") if isinstance(transcript, dict) else ""
        if raw_text:
            try:
                polished = _polish_transcript(raw_text, openai_key)
                if polished:
                    transcript["text"] = polished
                    transcript["segments"] = []
            except Exception as polish_err:
                print(f"[polish] GPT 加标点失败: {polish_err}")

        return jsonify({"transcript": transcript})

    except Exception as e:
        return jsonify({"error": f"转录失败: {e}"}), 500
    finally:
        for f in tmp_files:
            try:
                if os.path.exists(f):
                    os.unlink(f)
            except Exception:
                pass


def _extract_aweme_id(video_url: str) -> str:
    """从抖音视频URL中提取 aweme_id"""
    if not video_url:
        return ""
    m = re.search(r'/video/(\d+)', video_url)
    return m.group(1) if m else ""


def _refresh_video_urls(aweme_id: str, cookie: str) -> dict:
    """通过抖音 API 获取视频最新的音频/视频下载 URL"""
    try:
        import httpx

        params = {
            "device_platform": "webapp",
            "aid": "6383",
            "channel": "channel_pc_web",
            "aweme_id": aweme_id,
            "pc_client_type": "1",
            "version_code": "290100",
            "version_name": "29.1.0",
            "cookie_enabled": "true",
            "platform": "PC",
            "msToken": _gen_mstoken(),
        }
        param_str = "&".join(f"{k}={v}" for k, v in params.items())
        xb = _XBogus(_DY_UA)
        signed_url = xb.get_xbogus(param_str)
        full_url = f"https://www.douyin.com/aweme/v1/web/aweme/detail/?{signed_url}"

        headers = {
            "User-Agent": _DY_UA,
            "Referer": "https://www.douyin.com/",
            "Cookie": cookie,
            "Accept": "application/json, text/plain, */*",
        }

        with httpx.Client(timeout=15, follow_redirects=True) as client:
            resp = client.get(full_url, headers=headers)
            resp.raise_for_status()
            data = resp.json()

        item = data.get("aweme_detail", {})
        if not item:
            return {}

        # 提取音频URL
        music = item.get("music", {})
        play_url_list = (music.get("play_url") or {}).get("url_list", [])
        audio_url = play_url_list[0] if play_url_list else ""

        # 提取视频下载URL
        video_obj = item.get("video", {})
        video_download_url = ""
        download_addr = video_obj.get("download_addr", {})
        da_list = download_addr.get("url_list", [])
        if da_list:
            video_download_url = da_list[0]
        if not video_download_url:
            play_addr = video_obj.get("play_addr", {})
            pa_list = play_addr.get("url_list", [])
            if pa_list:
                video_download_url = pa_list[0]

        return {"audio_url": audio_url, "video_download_url": video_download_url}
    except Exception as e:
        # 不再静默吞掉异常，返回错误信息便于诊断
        return {"_error": str(e)}


def _download_url(url: str, dest_path: str, cookie: str = "", proxy: str = "",
                   max_bytes: int = 0):
    """下载URL到本地文件。先用 httpx，失败后降级 urllib。
    max_bytes>0 时用 Range 头限制下载量。
    """
    headers = {
        "User-Agent": _DY_UA,
        "Referer": "https://www.douyin.com/",
        "Accept": "*/*",
        "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
        "Accept-Encoding": "identity",
        "Connection": "keep-alive",
    }
    if cookie:
        headers["Cookie"] = cookie
    if max_bytes > 0:
        headers["Range"] = f"bytes=0-{max_bytes - 1}"

    last_err = None

    # 方式1: httpx（支持 HTTP/2，更好的 TLS 指纹）
    try:
        import httpx
        for attempt in range(2):
            try:
                with httpx.Client(
                    timeout=90,
                    follow_redirects=True,
                    http2=(attempt == 0),
                    proxy=proxy if proxy else None,
                ) as client:
                    with client.stream("GET", url, headers=headers) as resp:
                        if resp.status_code not in (200, 206):
                            resp.raise_for_status()
                        downloaded = 0
                        with open(dest_path, "wb") as f:
                            for chunk in resp.iter_bytes(8192):
                                f.write(chunk)
                                downloaded += len(chunk)
                                if max_bytes > 0 and downloaded >= max_bytes:
                                    break
                if os.path.getsize(dest_path) >= 1000:
                    return
            except Exception as e:
                last_err = e
                time.sleep(1)
    except ImportError:
        pass

    # 方式2: urllib（不同 TLS 指纹，作为降级方案）
    import ssl
    try:
        req = urllib.request.Request(url)
        for k, v in headers.items():
            req.add_header(k, v)
        ctx = ssl.create_default_context()
        resp = urllib.request.urlopen(req, timeout=90, context=ctx)
        downloaded = 0
        with open(dest_path, "wb") as f:
            while True:
                chunk = resp.read(8192)
                if not chunk:
                    break
                f.write(chunk)
                downloaded += len(chunk)
                if max_bytes > 0 and downloaded >= max_bytes:
                    break
        if os.path.getsize(dest_path) >= 1000:
            return
    except Exception as e:
        last_err = e

    raise RuntimeError(f"下载失败(重试3次): {last_err}")


# ─── 文稿加标点 ───


@app.route("/api/polish", methods=["POST"])
def polish():
    """给单个视频文稿添加标点符号"""
    data = request.json or {}
    cfg = _get_config(data)
    text = data.get("text", "").strip()
    openai_key = cfg["openai_api_key"]

    if not text:
        return jsonify({"error": "缺少文本"}), 400
    if not openai_key:
        return jsonify({"error": "缺少 OpenAI API Key"}), 400

    # 已有足够标点则跳过
    if _has_punctuation(text):
        return jsonify({"polished": text, "skipped": True})

    try:
        polished = _polish_transcript(text, openai_key)
        if polished:
            return jsonify({"polished": polished, "skipped": False})
        return jsonify({"polished": text, "skipped": True})
    except Exception as e:
        return jsonify({"error": f"加标点失败: {e}"}), 500


# ─── 生成 Word 文档 ───


@app.route("/api/generate-doc", methods=["POST"])
def generate_doc():
    data = request.json or {}
    videos = data.get("videos", [])
    creator_name = data.get("creator_name", "博主")

    if not videos:
        return jsonify({"error": "没有视频数据"}), 400

    try:
        doc_bytes = _generate_word_doc(videos, creator_name)
        safe_name = re.sub(r'[\\/:*?"<>|]', "_", creator_name)
        filename = f"{safe_name}_文字稿合集.docx"

        return send_file(
            BytesIO(doc_bytes),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename,
        )
    except Exception as e:
        return jsonify({"error": f"文档生成失败: {e}"}), 500


# ─── 博主 GPT 对话 ───


@app.route("/api/debug-video-fields", methods=["POST"])
def debug_video_fields():
    """诊断端点：返回第一个视频的完整字段结构，用于查找字幕相关字段"""
    data = request.json or {}
    cfg = _get_config(data)
    sec_uid = data.get("sec_uid", "").strip()
    cookie = cfg["cookie"]

    if not sec_uid:
        return jsonify({"error": "缺少 sec_uid"}), 400

    try:
        params = _build_base_params(sec_uid, 0, count=1)
        resp_data = _douyin_api_request(
            "https://www.douyin.com/aweme/v1/web/aweme/post/", params, cookie
        )
        aweme_list = resp_data.get("aweme_list", [])
        if not aweme_list:
            return jsonify({"error": "无视频数据"}), 404

        item = aweme_list[0]
        # 提取关键字段结构（不返回完整数据以避免太大）
        video_obj = item.get("video", {})
        result = {
            "video_keys": list(video_obj.keys()),
            # 音频流字段（DASH分离的纯音频）
            "audio": video_obj.get("audio"),
            "bit_rate_audio": video_obj.get("bit_rate_audio"),
            # play_addr_h264 结构
            "play_addr_h264": video_obj.get("play_addr_h264"),
            # download_addr 是否存在
            "has_download_addr": "download_addr" in video_obj,
            # 视频描述
            "desc": item.get("desc"),
        }
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/upload-corpus", methods=["POST"])
def upload_corpus():
    """上传文档文件，解析文本内容作为语料库"""
    if "file" not in request.files:
        return jsonify({"error": "未选择文件"}), 400

    file = request.files["file"]
    if not file.filename:
        return jsonify({"error": "文件名为空"}), 400

    avatar_name = request.form.get("avatar_name", "").strip()
    if not avatar_name:
        return jsonify({"error": "请输入分身名称"}), 400

    filename = file.filename.lower()
    try:
        if filename.endswith(".docx"):
            text = _parse_docx(file)
        elif filename.endswith(".txt") or filename.endswith(".md"):
            text = file.read().decode("utf-8", errors="ignore")
        elif filename.endswith(".pdf"):
            text = _parse_pdf(file)
        elif filename.endswith(".doc"):
            return jsonify({"error": "不支持 .doc 格式，请转为 .docx 后上传"}), 400
        else:
            # 尝试当纯文本读
            text = file.read().decode("utf-8", errors="ignore")

        text = text.strip()
        if not text:
            return jsonify({"error": "文件内容为空"}), 400

        return jsonify({
            "text": text,
            "char_count": len(text),
            "avatar_name": avatar_name,
        })
    except Exception as e:
        return jsonify({"error": f"文件解析失败: {e}"}), 500


def _parse_docx(file_obj) -> str:
    """解析 .docx 文件提取纯文本"""
    from docx import Document
    doc = Document(file_obj)
    parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    return "\n\n".join(parts)


def _parse_pdf(file_obj) -> str:
    """简易 PDF 文本提取（纯文本层）"""
    content = file_obj.read()
    # 提取 PDF 文本流中的文本（简易方案，不依赖额外库）
    text_parts = []
    # 查找 BT...ET 文本块中的 Tj/TJ 操作符
    for match in re.finditer(rb"\(([^)]*)\)\s*Tj", content):
        try:
            raw = match.group(1)
            # 处理 PDF 转义
            raw = raw.replace(b"\\(", b"(").replace(b"\\)", b")").replace(b"\\\\", b"\\")
            text_parts.append(raw.decode("utf-8", errors="ignore"))
        except Exception:
            pass
    if text_parts:
        return "\n".join(text_parts)
    # 如果上面没提取到，尝试暴力提取可读文本
    try:
        decoded = content.decode("utf-8", errors="ignore")
        # 过滤掉二进制垃圾，只保留中英文和标点
        clean = re.sub(r"[^\u4e00-\u9fff\u3000-\u303fa-zA-Z0-9\s.,;:!?，。；：！？、""''（）\-\n]", "", decoded)
        lines = [l.strip() for l in clean.split("\n") if len(l.strip()) > 5]
        if lines:
            return "\n".join(lines)
    except Exception:
        pass
    return ""


@app.route("/api/chat", methods=["POST"])
def chat():
    data = request.json or {}
    cfg = _get_config(data)
    message = data.get("message", "").strip()
    creator_name = data.get("creator_name", "博主")
    videos_context = data.get("videos_context", [])
    corpus_text = data.get("corpus_text", "").strip()
    corpus_instructions = data.get("corpus_instructions", "").strip()
    history = data.get("history", [])
    openai_key = cfg["openai_api_key"]
    proxy = cfg["proxy"]

    if not message:
        return jsonify({"error": "消息不能为空"}), 400
    if not openai_key:
        return jsonify({"error": "缺少 OpenAI API Key"}), 400

    try:
        if corpus_text:
            # 上传语料模式
            reply = _chat_with_corpus(
                message, creator_name, corpus_text, corpus_instructions,
                history, openai_key, proxy
            )
        else:
            # 博主视频模式
            reply = _chat_with_creator(
                message, creator_name, videos_context, history, openai_key, proxy
            )
        return jsonify({"response": reply})
    except Exception as e:
        return jsonify({"error": f"对话失败: {e}"}), 500


# ═══════════════════════════════════════════════════
# 内部实现函数
# ═══════════════════════════════════════════════════


def _resolve_douyin_input(user_input: str) -> str:
    """解析用户输入为抖音主页 URL"""
    user_input = user_input.strip()

    # 短链接
    if "v.douyin.com" in user_input or "douyin.com/share" in user_input:
        url_match = re.search(r"https?://[^\s]+", user_input)
        if url_match:
            short_url = url_match.group(0)
            try:
                req = urllib.request.Request(short_url, method="HEAD")
                req.add_header("User-Agent", "Mozilla/5.0")
                resp = urllib.request.urlopen(req, timeout=10)
                final_url = resp.url
                sec_uid_match = re.search(r"sec_uid=([^&]+)", final_url)
                if sec_uid_match:
                    sec_uid = urllib.parse.unquote(sec_uid_match.group(1))
                    return f"https://www.douyin.com/user/{sec_uid}"
                user_match = re.search(r"/user/([^?&]+)", final_url)
                if user_match:
                    return f"https://www.douyin.com/user/{user_match.group(1)}"
            except Exception:
                pass

    # 完整 URL
    if "douyin.com/user/" in user_input:
        user_match = re.search(r"douyin\.com/user/([^?&\s]+)", user_input)
        if user_match:
            return f"https://www.douyin.com/user/{user_match.group(1)}"

    # 纯 sec_uid
    if user_input.startswith("MS4wLjABAAAA"):
        return f"https://www.douyin.com/user/{user_input}"

    return f"https://www.douyin.com/user/{user_input}"


def _extract_sec_uid(profile_url: str) -> str:
    """从主页 URL 提取 sec_uid"""
    match = re.search(r"/user/([^?&\s]+)", profile_url)
    return match.group(1) if match else ""


def _match_keyword(video: dict, keywords: list[str]) -> bool:
    """检查视频是否匹配关键词"""
    title = video.get("title", "") or ""
    raw_title = video.get("raw_title", "") or ""
    return any(kw in title or kw in raw_title for kw in keywords)


# ─── XBogus 签名（纯 Python，无外部依赖）───


class _XBogus:
    """抖音 X-Bogus 签名算法（来自 f2 项目，Apache 2.0 协议）"""

    def __init__(self, user_agent: str = "") -> None:
        self.Array = [
            None, None, None, None, None, None, None, None, None, None, None, None, None, None, None, None,
            None, None, None, None, None, None, None, None, None, None, None, None, None, None, None, None,
            None, None, None, None, None, None, None, None, None, None, None, None, None, None, None, None,
            0, 1, 2, 3, 4, 5, 6, 7, 8, 9, None, None, None, None, None, None, None, None, None, None, None,
            None, None, None, None, None, None, None, None, None, None, None, None, None, None, None, None,
            None, None, None, None, None, None, None, None, None, None, None, None, 10, 11, 12, 13, 14, 15
        ]
        self.character = "Dkdpgh4ZKsQB80/Mfvw36XI1R25-WUAlEi7NLboqYTOPuzmFjJnryx9HVGcaStCe="
        self.ua_key = b"\x00\x01\x0c"
        self.user_agent = user_agent or (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
            "(KHTML, like Gecko) Chrome/130.0.0.0 Safari/537.36 Edg/130.0.0.0"
        )

    def _md5_str_to_array(self, md5_str):
        if isinstance(md5_str, str) and len(md5_str) > 32:
            return [ord(c) for c in md5_str]
        array, idx = [], 0
        while idx < len(md5_str):
            array.append((self.Array[ord(md5_str[idx])] << 4) | self.Array[ord(md5_str[idx + 1])])
            idx += 2
        return array

    def _md5(self, input_data):
        if isinstance(input_data, str):
            arr = self._md5_str_to_array(input_data)
        else:
            arr = input_data
        return hashlib.md5(bytes(arr)).hexdigest()

    def _rc4_encrypt(self, key, data):
        S = list(range(256))
        j = 0
        for i in range(256):
            j = (j + S[i] + key[i % len(key)]) % 256
            S[i], S[j] = S[j], S[i]
        i = j = 0
        out = bytearray()
        for byte in data:
            i = (i + 1) % 256
            j = (j + S[i]) % 256
            S[i], S[j] = S[j], S[i]
            out.append(byte ^ S[(S[i] + S[j]) % 256])
        return out

    def _calc(self, a1, a2, a3):
        x3 = ((a1 & 255) << 16) | ((a2 & 255) << 8) | a3
        return (self.character[(x3 & 16515072) >> 18] + self.character[(x3 & 258048) >> 12]
                + self.character[(x3 & 4032) >> 6] + self.character[x3 & 63])

    def get_xbogus(self, url_params):
        a1 = self._md5_str_to_array(self._md5(
            base64.b64encode(self._rc4_encrypt(self.ua_key, self.user_agent.encode("ISO-8859-1"))).decode("ISO-8859-1")
        ))
        a2 = self._md5_str_to_array(self._md5(self._md5_str_to_array("d41d8cd98f00b204e9800998ecf8427e")))
        up = self._md5_str_to_array(self._md5(self._md5_str_to_array(self._md5(url_params))))
        timer = int(time.time())
        ct = 536919696
        na = [64, 0, 1, 12, up[14], up[15], a2[14], a2[15], a1[14], a1[15],
              timer >> 24 & 255, timer >> 16 & 255, timer >> 8 & 255, timer & 255,
              ct >> 24 & 255, ct >> 16 & 255, ct >> 8 & 255, ct & 255]
        xor_r = na[0]
        for v in na[1:]:
            xor_r ^= int(v)
        na.append(xor_r)
        a3, a4 = [], []
        for idx in range(0, len(na), 2):
            a3.append(na[idx])
            if idx + 1 < len(na):
                a4.append(na[idx + 1])
        merge = a3 + a4
        # encoding_conversion 参数顺序: a,b,c,e,d,t,f,r,n,o,i,_,x,u,s,l,v,h,p
        # y = [a, int(i), b, _, c, x, e, u, d, s, t, l, f, v, r, h, n, p, o]
        m = merge
        y = [m[0], int(m[10]), m[1], m[11], m[2], m[12], m[3], m[13],
             m[4], m[14], m[5], m[15], m[6], m[16], m[7], m[17],
             m[8], m[18] if len(m) > 18 else 0, m[9]]
        garbled = chr(2) + chr(255) + self._rc4_encrypt(
            "ÿ".encode("ISO-8859-1"), bytes(y[:19]).decode("ISO-8859-1").encode("ISO-8859-1")
        ).decode("ISO-8859-1")
        xb_ = ""
        idx = 0
        while idx < len(garbled):
            xb_ += self._calc(ord(garbled[idx]), ord(garbled[idx + 1]), ord(garbled[idx + 2]))
            idx += 3
        return f"{url_params}&X-Bogus={xb_}"


# ─── 直接 HTTP 抖音 API（替代 f2）───

_DY_UA = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
    "(KHTML, like Gecko) Chrome/130.0.0.0 Safari/537.36 Edg/130.0.0.0"
)


def _gen_mstoken() -> str:
    """生成随机 msToken"""
    chars = string.ascii_letters + string.digits + "+/"
    return "".join(random.choices(chars, k=126)) + "=="


def _build_base_params(sec_uid: str, max_cursor: int = 0, count: int = 20) -> dict:
    """构建抖音 API 请求参数"""
    return {
        "device_platform": "webapp",
        "aid": "6383",
        "channel": "channel_pc_web",
        "sec_user_id": sec_uid,
        "max_cursor": str(max_cursor),
        "locate_query": "false",
        "show_live_replay_strategy": "1",
        "need_time_list": "1",
        "time_list_query": "0",
        "whale_cut_token": "",
        "cut_version": "1",
        "count": str(count),
        "publish_video_strategy_type": "2",
        "pc_client_type": "1",
        "version_code": "290100",
        "version_name": "29.1.0",
        "cookie_enabled": "true",
        "screen_width": "1920",
        "screen_height": "1080",
        "browser_language": "zh-CN",
        "browser_platform": "Win32",
        "browser_name": "Edge",
        "browser_version": "130.0.0.0",
        "browser_online": "true",
        "engine_name": "Blink",
        "engine_version": "130.0.0.0",
        "os_name": "Windows",
        "os_version": "10",
        "cpu_core_num": "12",
        "device_memory": "8",
        "platform": "PC",
        "downlink": "10",
        "effective_type": "4g",
        "round_trip_time": "100",
        "msToken": _gen_mstoken(),
    }


def _douyin_api_request(endpoint: str, params: dict, cookie: str) -> dict:
    """发送带签名的抖音 API 请求"""
    import httpx

    param_str = "&".join(f"{k}={v}" for k, v in params.items())
    xb = _XBogus(_DY_UA)
    signed_url = xb.get_xbogus(param_str)
    full_url = f"{endpoint}?{signed_url}"

    headers = {
        "User-Agent": _DY_UA,
        "Referer": "https://www.douyin.com/",
        "Cookie": cookie,
        "Accept": "application/json, text/plain, */*",
    }

    with httpx.Client(timeout=30, follow_redirects=True) as client:
        resp = client.get(full_url, headers=headers)
        resp.raise_for_status()
        data = resp.json()

    # 检查抖音 API 业务错误码
    status_code = data.get("status_code")
    if status_code and status_code != 0:
        msg = data.get("status_msg", "")
        raise RuntimeError(f"抖音 API 错误 (code={status_code}): {msg or '请检查 Cookie 是否有效'}")

    return data


def _fetch_user_profile_direct(sec_uid: str, cookie: str) -> dict:
    """直接调用抖音 API 获取用户信息"""
    params = _build_base_params(sec_uid)
    return _douyin_api_request("https://www.douyin.com/aweme/v1/web/user/profile/other/", params, cookie)


def _fetch_videos_direct(
    sec_uid: str, cookie: str, max_videos: int = 0, keyword: str = ""
) -> tuple[list[dict], str]:
    """直接调用抖音 API 获取视频列表（替代 f2）"""
    # 获取用户信息
    profile_data = _fetch_user_profile_direct(sec_uid, cookie)
    user = profile_data.get("user", {})
    creator_name = user.get("nickname", "")

    keywords = keyword.split() if keyword else None
    all_videos = []
    max_cursor = 0
    max_count = max_videos if max_videos > 0 else 9999

    while len(all_videos) < max_count:
        params = _build_base_params(sec_uid, max_cursor, count=35)
        data = _douyin_api_request(
            "https://www.douyin.com/aweme/v1/web/aweme/post/", params, cookie
        )

        aweme_list = data.get("aweme_list", [])
        if not aweme_list:
            break

        for item in aweme_list:
            vid = item.get("aweme_id", "")
            desc = item.get("desc", "无标题")
            title = re.sub(r"#\S+", "", desc).strip() or "无标题"
            duration = item.get("video", {}).get("duration", 0)
            if isinstance(duration, (int, float)) and duration > 10000:
                duration = duration // 1000
            # 视频下载URL：优先 download_addr（完整MP4含音频），play_addr是fMP4仅含视频
            video_obj = item.get("video", {})
            video_download_url = ""
            # 方式1: download_addr（完整MP4，含音频轨道）
            download_addr = video_obj.get("download_addr", {})
            da_list = download_addr.get("url_list", [])
            if da_list:
                video_download_url = da_list[0]
            # 方式2: play_addr（DASH fMP4，可能仅含视频轨道）
            if not video_download_url:
                play_addr = video_obj.get("play_addr", {})
                pa_list = play_addr.get("url_list", [])
                if pa_list:
                    video_download_url = pa_list[0]
            # 方式3: 最低码率（备选）
            if not video_download_url:
                bit_rates = video_obj.get("bit_rate", [])
                if bit_rates:
                    lowest = min(bit_rates, key=lambda b: b.get("bit_rate", float("inf")))
                    br_addr = lowest.get("play_addr", {})
                    br_list = br_addr.get("url_list", [])
                    video_download_url = br_list[0] if br_list else ""
            # 背景音乐URL（最后备选）
            music = item.get("music", {})
            play_url_list = (music.get("play_url") or {}).get("url_list", [])
            audio_url = play_url_list[0] if play_url_list else ""
            author = (item.get("author") or {}).get("nickname", creator_name)
            create_time = item.get("create_time", "")

            video = {
                "id": str(vid),
                "title": title,
                "raw_title": desc,
                "url": f"https://www.douyin.com/video/{vid}",
                "create_time": create_time,
                "duration": duration,
                "author": author,
                "video_download_url": video_download_url,
                "audio_url": audio_url,
                "creator_name": creator_name,
            }

            if keywords:
                if _match_keyword(video, keywords):
                    all_videos.append(video)
            else:
                all_videos.append(video)

            if len(all_videos) >= max_count:
                break

        has_more = data.get("has_more", False)
        max_cursor = data.get("max_cursor", 0)
        if not has_more or not max_cursor:
            break

        # 分页间短延迟，防止触发反爬
        time.sleep(0.3)

    return all_videos, creator_name


# ─── Apify 视频获取 ───


def _apify_fetch_videos(
    apify_token: str, profile_url: str, max_videos: int = 200
) -> list[dict]:
    """通过 Apify 获取视频列表"""
    from apify_client import ApifyClient

    client = ApifyClient(apify_token)

    actors = [
        {
            "id": "natanielsantos/douyin-scraper",
            "input": {
                "profileUrls": [profile_url],
                "searchTermsOrHashtags": [],
                "postUrls": [],
                "maxItemsPerUrl": max_videos,
                "profileSortFilter": "latest",
            },
        },
    ]

    for actor in actors:
        try:
            run = client.actor(actor["id"]).call(run_input=actor["input"])
            items = list(client.dataset(run["defaultDatasetId"]).iterate_items())
            if items:
                return _normalize_video_list(items)
        except Exception:
            continue

    return []


def _normalize_video_list(items: list) -> list[dict]:
    """标准化视频列表"""
    videos = []
    seen_ids = set()

    for idx, item in enumerate(items):
        vid = str(item.get("id", item.get("aweme_id", f"_no_id_{idx}")))
        if vid in seen_ids:
            continue
        seen_ids.add(vid)

        raw_title = (
            item.get("text")
            or item.get("desc")
            or item.get("title")
            or item.get("description")
            or "无标题"
        )
        title = re.sub(r"#\S+", "", raw_title).strip() or "无标题"

        author_meta = item.get("authorMeta") or {}
        author = author_meta.get("name") or item.get("author") or ""

        stats = item.get("statistics") or {}
        digg_count = stats.get("diggCount") or item.get("digg_count") or 0

        video_meta = item.get("videoMeta") or {}
        duration = video_meta.get("duration") or item.get("duration") or 0
        if duration > 10000:
            duration = duration // 1000

        music_meta = item.get("musicMeta") or {}
        audio_url = music_meta.get("playUrl") or ""

        videos.append({
            "id": vid,
            "title": title,
            "raw_title": raw_title,
            "url": item.get("url", ""),
            "create_time": item.get("createTime", item.get("create_time", "")),
            "duration": duration,
            "digg_count": digg_count,
            "author": author,
            "audio_url": audio_url,
        })

    videos.sort(key=lambda x: str(x.get("create_time", "")), reverse=True)
    return videos


# ─── Whisper 转录 ───


def _parse_boxes(data: bytes) -> list[tuple]:
    """解析 MP4 box 结构，返回 [(offset, size, type_bytes), ...]"""
    import struct
    boxes = []
    pos = 0
    while pos < len(data):
        if pos + 8 > len(data):
            break
        size = struct.unpack(">I", data[pos:pos + 4])[0]
        box_type = data[pos + 4:pos + 8]
        if size == 0:
            size = len(data) - pos
        elif size == 1:
            if pos + 16 > len(data):
                break
            size = struct.unpack(">Q", data[pos + 8:pos + 16])[0]
        if size < 8 or pos + size > len(data):
            break
        boxes.append((pos, size, box_type))
        pos += size
    return boxes


def _get_trak_handler(trak_content: bytes) -> str:
    """从 trak 内容中找到 handler_type（遍历 mdia → hdlr）。
    返回 handler_type 字符串，如 'soun'、'vide'，或空字符串。
    """
    # 扫描 trak 子 box 找 mdia
    for _, _, bt in _parse_boxes(trak_content):
        pass  # 只是验证能解析
    for bpos, bsize, bt in _parse_boxes(trak_content):
        if bt == b"mdia":
            mdia_content = trak_content[bpos + 8:bpos + bsize]
            # 扫描 mdia 子 box 找 hdlr
            for mpos, msize, mt in _parse_boxes(mdia_content):
                if mt == b"hdlr":
                    # hdlr FullBox: 4 version+flags + 4 pre_defined + 4 handler_type
                    hdlr_content = mdia_content[mpos + 8:mpos + msize]
                    if len(hdlr_content) >= 12:
                        return hdlr_content[8:12].decode("ascii", errors="replace")
            break
    return ""


def _extract_audio(video_path: str) -> str:
    """用 ffmpeg 从视频中提取纯音频（AAC/M4A），文件大小大幅缩小。"""
    import subprocess

    audio_path = video_path.rsplit(".", 1)[0] + ".m4a"

    ffmpeg_bin = _get_ffmpeg_bin()

    cmd = [
        ffmpeg_bin,
        "-i", video_path,
        "-vn",              # 去掉视频
        "-acodec", "copy",  # 音频直接复制，不重编码
        "-y",               # 覆盖
        audio_path,
    ]
    result = subprocess.run(cmd, capture_output=True, timeout=60)
    if result.returncode != 0:
        stderr = result.stderr.decode("utf-8", errors="replace")[-500:]
        raise RuntimeError(f"ffmpeg 提取音频失败: {stderr}")

    if not os.path.exists(audio_path) or os.path.getsize(audio_path) < 1000:
        raise RuntimeError("ffmpeg 输出文件为空")

    return audio_path


def _get_ffmpeg_bin() -> str:
    """获取 ffmpeg 可执行文件路径"""
    try:
        import imageio_ffmpeg
        return imageio_ffmpeg.get_ffmpeg_exe()
    except ImportError:
        return "ffmpeg"


def _compress_audio(input_path: str, tmp_files: list) -> str:
    """用 ffmpeg 将音频压缩为低码率 MP3（目标 <25MB）"""
    import subprocess

    mp3_path = input_path.rsplit(".", 1)[0] + "_compressed.mp3"
    ffmpeg_bin = _get_ffmpeg_bin()

    cmd = [
        ffmpeg_bin,
        "-i", input_path,
        "-vn",
        "-acodec", "libmp3lame",
        "-ab", "32k",       # 极低码率，保证文件够小
        "-ar", "16000",     # 16kHz 采样率（Whisper 足够）
        "-ac", "1",         # 单声道
        "-y",
        mp3_path,
    ]
    result = subprocess.run(cmd, capture_output=True, timeout=120)
    if result.returncode != 0:
        stderr = result.stderr.decode("utf-8", errors="replace")[-500:]
        raise RuntimeError(f"ffmpeg 压缩失败: {stderr}")

    if not os.path.exists(mp3_path) or os.path.getsize(mp3_path) < 1000:
        raise RuntimeError("ffmpeg 压缩输出为空")

    return mp3_path


def _split_and_transcribe(file_path: str, api_key: str, proxy: str, tmp_files: list) -> str:
    """大文件分段：每段 24MB，分别转录后拼接文本"""
    max_chunk = 24 * 1024 * 1024
    file_size = os.path.getsize(file_path)
    all_text = []

    with open(file_path, "rb") as f:
        chunk_idx = 0
        while True:
            data = f.read(max_chunk)
            if not data:
                break
            chunk_path = file_path.rsplit(".", 1)[0] + f"_chunk{chunk_idx}.mp4"
            with open(chunk_path, "wb") as cf:
                cf.write(data)
            tmp_files.append(chunk_path)

            try:
                result = _call_whisper(chunk_path, api_key, proxy)
                text = result.get("text", "") if isinstance(result, dict) else ""
                if text:
                    all_text.append(text)
            except Exception:
                pass  # 跳过失败的段

            chunk_idx += 1

    return "".join(all_text)


def _has_punctuation(text: str) -> bool:
    """检查中文文本是否包含足够标点"""
    if not text or len(text) < 20:
        return True
    puncts = sum(1 for c in text if c in "，。！？、；：""''（）…—")
    ratio = puncts / len(text)
    return ratio > 0.01  # 至少 1% 的字符是标点


def _polish_transcript(raw_text: str, api_key: str) -> str:
    """用 GPT 给转录文稿添加标点符号，整理为标准文稿"""
    import ssl

    if not raw_text or len(raw_text.strip()) < 10:
        return ""

    # 长文本分块处理，每块不超过 2000 字
    chunks = []
    text = raw_text.strip()
    while len(text) > 2000:
        chunks.append(text[:2000])
        text = text[2000:]
    chunks.append(text)

    polished_parts = []
    for chunk in chunks:
        payload = json.dumps({
            "model": "gpt-4.1-mini",
            "max_tokens": 8192,
            "messages": [
                {
                    "role": "system",
                    "content": (
                        "你是一个中文文稿整理专家。请将以下语音转录文本整理为标准文稿：\n"
                        "1. 添加正确的中文标点符号（句号、逗号、问号、感叹号、顿号等）\n"
                        "2. 去除口水词（嗯、啊、呃、那个、就是说等）\n"
                        "3. 不要改变原文的意思和用词\n"
                        "4. 不要添加任何解释或注释\n"
                        "5. 直接输出整理后的文稿，不要有任何前缀说明"
                    ),
                },
                {"role": "user", "content": chunk},
            ],
        }).encode()

        req = urllib.request.Request(
            "https://api.openai.com/v1/chat/completions",
            data=payload,
            method="POST",
        )
        req.add_header("Authorization", f"Bearer {api_key.strip()}")
        req.add_header("Content-Type", "application/json")

        ctx = ssl.create_default_context()
        resp = urllib.request.urlopen(req, timeout=60, context=ctx)
        result = json.loads(resp.read().decode())

        choices = result.get("choices", [])
        if choices:
            part = choices[0].get("message", {}).get("content", "").strip()
            if part:
                polished_parts.append(part)
            else:
                polished_parts.append(chunk)
        else:
            polished_parts.append(chunk)

    return "".join(polished_parts)


def _call_whisper(audio_path: str, api_key: str, proxy: str = "") -> dict:
    """调用 OpenAI Whisper API（直接 HTTP，不依赖 openai/httpx SDK）"""
    import mimetypes
    import ssl
    import urllib.error

    boundary = "----WhisperBoundary" + "".join(random.choices(string.ascii_letters, k=16))
    filename = os.path.basename(audio_path)
    content_type = mimetypes.guess_type(filename)[0] or "application/octet-stream"

    # 构建 multipart body
    body_parts = []
    # model
    body_parts.append(f"--{boundary}\r\nContent-Disposition: form-data; name=\"model\"\r\n\r\nwhisper-1".encode())
    # language
    body_parts.append(f"--{boundary}\r\nContent-Disposition: form-data; name=\"language\"\r\n\r\nzh".encode())
    # response_format
    body_parts.append(f"--{boundary}\r\nContent-Disposition: form-data; name=\"response_format\"\r\n\r\nverbose_json".encode())
    # timestamp_granularities[]
    body_parts.append(f"--{boundary}\r\nContent-Disposition: form-data; name=\"timestamp_granularities[]\"\r\n\r\nsegment".encode())
    # prompt
    body_parts.append(f"--{boundary}\r\nContent-Disposition: form-data; name=\"prompt\"\r\n\r\n以下是普通话的句子，包含标点符号。".encode())
    # file
    with open(audio_path, "rb") as f:
        file_data = f.read()
    body_parts.append(
        f"--{boundary}\r\nContent-Disposition: form-data; name=\"file\"; filename=\"{filename}\"\r\nContent-Type: {content_type}\r\n\r\n".encode()
        + file_data
    )
    body_parts.append(f"--{boundary}--\r\n".encode())
    body = b"\r\n".join(body_parts)

    req = urllib.request.Request(
        "https://api.openai.com/v1/audio/transcriptions",
        data=body,
        method="POST",
    )
    req.add_header("Authorization", f"Bearer {api_key.strip()}")
    req.add_header("Content-Type", f"multipart/form-data; boundary={boundary}")

    ctx = ssl.create_default_context()
    try:
        resp = urllib.request.urlopen(req, timeout=45, context=ctx)
    except urllib.error.HTTPError as http_err:
        # 读取错误响应体，获取具体原因
        err_body = ""
        try:
            err_body = http_err.read().decode("utf-8", errors="replace")[:500]
        except Exception:
            pass
        raise RuntimeError(f"Whisper HTTP {http_err.code}: {err_body}") from http_err
    result = json.loads(resp.read().decode())

    segments = []
    for seg in result.get("segments", []):
        segments.append({
            "start": round(seg.get("start", 0), 2),
            "end": round(seg.get("end", 0), 2),
            "text": seg.get("text", "").strip(),
        })

    return {
        "text": result.get("text", "").strip(),
        "segments": segments,
        "language": result.get("language", "zh"),
    }


# ─── Word 文档生成 ───


def _remove_filler_words(text: str) -> str:
    """去除中文口水词/填充词"""
    # 常见口水词列表
    fillers = [
        r"(?<=[。！？\s])嗯+[，、。]?",
        r"(?<=[。！？\s])啊+[，、。]?",
        r"(?<=[。！？\s])呃+[，、。]?",
        r"(?<=[。！？\s])额+[，、。]?",
        r"(?<=[\s，。])那个[，、]?(?=\S)",
        r"(?<=[\s，。])就是说[，、]?",
        r"(?<=[\s，。])然后的话[，、]?",
        r"(?<=[\s，。])对不对[，、。]?",
        r"(?<=[\s，。])是不是[，、。]?",
        r"(?<=[\s，。])你知道吧[，、。]?",
        r"(?<=[\s，。])就是嘛[，、。]?",
        r"(?<=[\s，。])怎么说呢[，、。]?",
    ]
    # 独立出现的口水词（整句或句首）
    standalone = [
        r"^嗯+[，、。]?\s*",
        r"^啊+[，、。]?\s*",
        r"^呃+[，、。]?\s*",
        r"^额+[，、。]?\s*",
    ]
    for pattern in fillers:
        text = re.sub(pattern, "", text)
    for pattern in standalone:
        text = re.sub(pattern, "", text)
    # 清理多余空格和重复标点
    text = re.sub(r"\s{2,}", " ", text)
    text = re.sub(r"[，、]{2,}", "，", text)
    text = re.sub(r"[。]{2,}", "。", text)
    return text.strip()


def _set_run_font(run, font_name: str = "微软雅黑", size_pt: int = 12):
    """设置 run 的中文字体（含 east_asia fallback）"""
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = font_name
    run.font.size = Pt(size_pt)
    # 设置东亚字体
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from lxml import etree
        rfonts = etree.SubElement(rpr, qn("w:rFonts"))
    rfonts.set(qn("w:eastAsia"), font_name)


def _add_bookmark(paragraph, bookmark_name: str):
    """在段落中添加书签"""
    from docx.oxml.ns import qn
    from lxml import etree

    tag_id = str(abs(hash(bookmark_name)) % 1000000)
    start = etree.SubElement(paragraph._element, qn("w:bookmarkStart"))
    start.set(qn("w:id"), tag_id)
    start.set(qn("w:name"), bookmark_name)
    end = etree.SubElement(paragraph._element, qn("w:bookmarkEnd"))
    end.set(qn("w:id"), tag_id)


def _add_hyperlink(paragraph, bookmark_name: str, text: str, font_size: int = 12):
    """在段落中添加指向书签的超链接"""
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
    from lxml import etree

    hyperlink = etree.SubElement(paragraph._element, qn("w:hyperlink"))
    hyperlink.set(qn("w:anchor"), bookmark_name)

    run_elem = etree.SubElement(hyperlink, qn("w:r"))
    rpr = etree.SubElement(run_elem, qn("w:rPr"))

    # 蓝色 + 下划线
    color = etree.SubElement(rpr, qn("w:color"))
    color.set(qn("w:val"), "0071E3")
    underline = etree.SubElement(rpr, qn("w:u"))
    underline.set(qn("w:val"), "single")
    sz = etree.SubElement(rpr, qn("w:sz"))
    sz.set(qn("w:val"), str(font_size * 2))
    sz_cs = etree.SubElement(rpr, qn("w:szCs"))
    sz_cs.set(qn("w:val"), str(font_size * 2))
    # 字体
    rfonts = etree.SubElement(rpr, qn("w:rFonts"))
    rfonts.set(qn("w:eastAsia"), "微软雅黑")

    text_elem = etree.SubElement(run_elem, qn("w:t"))
    text_elem.text = text
    text_elem.set(qn("xml:space"), "preserve")


def _generate_word_doc(videos: list[dict], creator_name: str) -> bytes:
    """生成 Word 文档，返回字节"""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()

    # 默认样式 — 使用中文字体
    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.font.name = "微软雅黑"
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from lxml import etree
        rfonts = etree.SubElement(rpr, qn("w:rFonts"))
    rfonts.set(qn("w:eastAsia"), "微软雅黑")
    style.paragraph_format.line_spacing = 1.5

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # ═══ 封面 ═══
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(creator_name)
    _set_run_font(run, "微软雅黑", 36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("视频文字稿合集")
    _set_run_font(run, "微软雅黑", 20)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    total = len(videos)
    transcribed = sum(1 for v in videos if v.get("transcript"))
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"共 {total} 个视频 · 已转录 {transcribed} 个")
    _set_run_font(run, "微软雅黑", 12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(f"生成日期：{datetime.now().strftime('%Y年%m月%d日')}")
    _set_run_font(run, "微软雅黑", 11)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.add_page_break()

    # ═══ 目录（可点击跳转）═══
    toc_heading = doc.add_heading("目 录", level=1)
    chapter_num = 0
    for v in videos:
        if v.get("transcript"):
            chapter_num += 1
            t = re.sub(r"#\S+", "", v.get("title", f"视频 {chapter_num}")).strip()[:60]
            bookmark_name = f"chapter_{chapter_num}"
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            _add_hyperlink(p, bookmark_name, f"{chapter_num}. {t}", font_size=12)

    doc.add_page_break()

    # ═══ 正文 ═══
    chapter_num = 0
    for v in videos:
        transcript = v.get("transcript")
        if not transcript:
            continue
        chapter_num += 1
        t = re.sub(r"#\S+", "", v.get("title", f"视频 {chapter_num}")).strip()[:80]
        bookmark_name = f"chapter_{chapter_num}"

        heading = doc.add_heading(f"{chapter_num}. {t}", level=2)
        _add_bookmark(heading, bookmark_name)

        # 元信息
        meta_parts = []
        ct = v.get("create_time", "")
        if ct:
            if isinstance(ct, (int, float)):
                try:
                    ct = datetime.fromtimestamp(ct).strftime("%Y-%m-%d")
                except Exception:
                    ct = ""
            if ct:
                meta_parts.append(f"发布时间：{str(ct)[:10]}")
        dur = v.get("duration", 0)
        if isinstance(dur, (int, float)) and dur > 0:
            m, s = divmod(int(dur), 60)
            meta_parts.append(f"时长：{m}:{s:02d}")
        if meta_parts:
            mp = doc.add_paragraph(" | ".join(meta_parts))
            for r in mp.runs:
                _set_run_font(r, "微软雅黑", 9)
                r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # 文字稿（去除口水词）
        text = ""
        segments = []
        if isinstance(transcript, dict):
            text = transcript.get("text", "")
            segments = transcript.get("segments", [])
        elif isinstance(transcript, str):
            text = transcript

        # 纯文稿：去除口水词，不带时间戳，标准段落格式
        if segments and len(segments) > 1:
            # 将所有 segment 文本合并为完整段落
            full_text = "".join(seg.get("text", "").strip() for seg in segments)
            cleaned = _remove_filler_words(full_text)
        elif text:
            cleaned = _remove_filler_words(text)
        else:
            cleaned = ""

        if cleaned:
            cp = doc.add_paragraph()
            cp.paragraph_format.line_spacing = 1.8
            cp.paragraph_format.first_line_indent = Pt(24)
            run = cp.add_run(cleaned)
            _set_run_font(run, "微软雅黑", 12)

        # 分隔线
        sep = doc.add_paragraph()
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sep.add_run("─" * 30)
        _set_run_font(run, "微软雅黑", 10)
        run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)

    # 输出字节
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _format_ts(seconds: float) -> str:
    seconds = max(0, int(seconds))
    m, s = divmod(seconds, 60)
    if m >= 60:
        h, m = divmod(m, 60)
        return f"{h}:{m:02d}:{s:02d}"
    return f"{m}:{s:02d}"


# ─── GPT 对话 ───


def _chat_with_creator(
    message: str,
    creator_name: str,
    videos_context: list[dict],
    history: list[dict],
    api_key: str,
    proxy: str = "",
) -> str:
    """与博主 GPT 对话（直接 HTTP，不依赖 httpx）"""

    # 构建博主资料
    profile_parts = []
    for v in videos_context[:50]:
        t = v.get("transcript")
        if t:
            text = t.get("text", "") if isinstance(t, dict) else str(t)
            if text:
                title = v.get("title", "")
                profile_parts.append(f"【{title}】\n{text[:2000]}")

    profile = "\n\n---\n\n".join(profile_parts) if profile_parts else "暂无内容"

    # 搜索相关内容
    query_words = set()
    phrases = re.split(r"[，。？！、；：\s,.\?!;:\n]+", message)
    for p in phrases:
        p = p.strip()
        if 2 <= len(p) <= 6:
            query_words.add(p)
        for n in (2, 3):
            for i in range(len(p) - n + 1):
                query_words.add(p[i : i + n])

    context_parts = []
    if query_words:
        scored = []
        for v in videos_context:
            t = v.get("transcript")
            if not t:
                continue
            text = t.get("text", "") if isinstance(t, dict) else str(t)
            score = sum(1 for w in query_words if w in text)
            if score > 0:
                scored.append((score, v.get("title", ""), text[:2000]))
        scored.sort(key=lambda x: x[0], reverse=True)
        for _, title, text in scored[:5]:
            context_parts.append(f"【{title}】\n{text}")

    context = "\n\n---\n\n".join(context_parts) if context_parts else "（无特别相关内容）"

    system_prompt = f"""你是抖音博主「{creator_name}」的 AI 分身。模仿这位博主的思维方式、说话风格来回答。

## 博主视频内容样本
{profile[:100000]}

## 当前对话相关参考
{context[:30000]}

## 规则
1. 用第一人称，保持博主风格
2. 优先用博主表达过的观点
3. 不说"根据视频"这样的元叙述
4. 自然口语化"""

    messages = [{"role": "system", "content": system_prompt}]
    messages.extend(history[-20:])
    messages.append({"role": "user", "content": message})

    # 直接 HTTP 调用 OpenAI Chat API（不依赖 httpx）
    import ssl
    payload = json.dumps({
        "model": "gpt-4.1",
        "max_tokens": 2048,
        "messages": messages,
    }).encode()

    req = urllib.request.Request(
        "https://api.openai.com/v1/chat/completions",
        data=payload,
        method="POST",
    )
    req.add_header("Authorization", f"Bearer {api_key.strip()}")
    req.add_header("Content-Type", "application/json")

    ctx = ssl.create_default_context()
    resp = urllib.request.urlopen(req, timeout=120, context=ctx)
    result = json.loads(resp.read().decode())

    choices = result.get("choices", [])
    if not choices:
        return "对话出错: 空响应"

    return choices[0].get("message", {}).get("content", "")


def _chat_with_corpus(
    message: str,
    avatar_name: str,
    corpus_text: str,
    corpus_instructions: str,
    history: list[dict],
    api_key: str,
    proxy: str = "",
) -> str:
    """基于上传语料的 GPT 对话"""

    # 搜索相关段落
    query_words = set()
    phrases = re.split(r"[，。？！、；：\s,.\?!;:\n]+", message)
    for p in phrases:
        p = p.strip()
        if 2 <= len(p) <= 6:
            query_words.add(p)
        for n in (2, 3):
            for i in range(len(p) - n + 1):
                query_words.add(p[i : i + n])

    # 按段落切分语料，找最相关的段落
    paragraphs = [p.strip() for p in corpus_text.split("\n\n") if p.strip()]
    if not paragraphs:
        paragraphs = [p.strip() for p in corpus_text.split("\n") if p.strip()]

    context_parts = []
    if query_words and paragraphs:
        scored = []
        for para in paragraphs:
            score = sum(1 for w in query_words if w in para)
            if score > 0:
                scored.append((score, para[:3000]))
        scored.sort(key=lambda x: x[0], reverse=True)
        for _, text in scored[:8]:
            context_parts.append(text)

    context = "\n\n---\n\n".join(context_parts) if context_parts else "（无特别相关段落）"

    instructions_block = ""
    if corpus_instructions:
        instructions_block = f"\n\n## 用户自定义要求\n{corpus_instructions}"

    system_prompt = f"""你是「{avatar_name}」，一个基于以下语料库内容的 AI 助手。根据语料库中的知识来回答问题。

## 语料库内容
{corpus_text[:120000]}

## 当前对话相关段落
{context[:30000]}
{instructions_block}

## 规则
1. 优先基于语料库中的内容回答
2. 如果语料库中没有相关信息，如实说明
3. 自然、专业地回答"""

    messages = [{"role": "system", "content": system_prompt}]
    messages.extend(history[-20:])
    messages.append({"role": "user", "content": message})

    import ssl
    payload = json.dumps({
        "model": "gpt-4.1",
        "max_tokens": 2048,
        "messages": messages,
    }).encode()

    req = urllib.request.Request(
        "https://api.openai.com/v1/chat/completions",
        data=payload,
        method="POST",
    )
    req.add_header("Authorization", f"Bearer {api_key.strip()}")
    req.add_header("Content-Type", "application/json")

    ctx = ssl.create_default_context()
    resp = urllib.request.urlopen(req, timeout=120, context=ctx)
    result = json.loads(resp.read().decode())

    choices = result.get("choices", [])
    if not choices:
        return "对话出错: 空响应"

    return choices[0].get("message", {}).get("content", "")
