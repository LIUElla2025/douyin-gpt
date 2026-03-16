"""Word 文档生成模块（docx 延迟加载，未安装时 app 仍能启动）"""

import re
from datetime import datetime
from pathlib import Path
from config import OUTPUT_DIR, sanitize_id


def _clean_title(title: str) -> str:
    """去掉标题中的 #话题标签"""
    return re.sub(r'#\S+', '', title).strip() or "无标题"


def _set_run_font(run, font_name, size_pt, color_rgb=None, bold=False, italic=False):
    """统一设置 run 的字体、字号、颜色"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    run.font.size = Pt(size_pt)
    run.font.name = font_name
    run.bold = bold
    run.italic = italic

    # 设置中文字体
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn('w:rFonts'))
    if fonts is None:
        fonts = OxmlElement('w:rFonts')
        rpr.append(fonts)
    fonts.set(qn('w:eastAsia'), font_name)
    fonts.set(qn('w:ascii'), font_name)
    fonts.set(qn('w:hAnsi'), font_name)

    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)


def _set_paragraph_spacing(paragraph, before_pt=0, after_pt=0, line_spacing=1.5):
    """设置段落间距和行距"""
    from docx.shared import Pt
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before_pt)
    fmt.space_after = Pt(after_pt)
    fmt.line_spacing = line_spacing


def _get_chinese_font():
    """根据系统选择中文字体"""
    import platform
    system = platform.system()
    if system == 'Darwin':
        return '苹方-简'
    elif system == 'Windows':
        return '微软雅黑'
    return 'Noto Sans CJK SC'


def _add_bookmark(paragraph, bookmark_name):
    """给段落添加书签"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), bookmark_name)
    bookmark_start.set(qn('w:name'), bookmark_name)
    paragraph._element.insert(0, bookmark_start)

    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), bookmark_name)
    paragraph._element.append(bookmark_end)


def _add_hyperlink_to_bookmark(paragraph, bookmark_name, text, font_name, size_pt=10.5, color=(0x33, 0x33, 0x33)):
    """添加指向书签的内部超链接"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)

    run_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # 字号
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size_pt * 2)))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(size_pt * 2)))
    rPr.append(szCs)

    # 字体
    fonts = OxmlElement('w:rFonts')
    fonts.set(qn('w:eastAsia'), font_name)
    fonts.set(qn('w:ascii'), font_name)
    fonts.set(qn('w:hAnsi'), font_name)
    rPr.append(fonts)

    # 颜色
    color_elem = OxmlElement('w:color')
    color_elem.set(qn('w:val'), '{:02X}{:02X}{:02X}'.format(*color))
    rPr.append(color_elem)

    run_elem.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.set(qn('xml:space'), 'preserve')
    text_elem.text = text
    run_elem.append(text_elem)
    hyperlink.append(run_elem)

    paragraph._element.append(hyperlink)


def generate_word_doc(
    videos: list[dict],
    creator_name: str,
    douyin_id: str,
) -> Path:
    """生成包含所有文字稿的 Word 文档"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx 未安装。请运行: pip install python-docx")

    font_name = _get_chinese_font()
    doc = Document()

    # 设置默认样式
    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.font.name = font_name
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        rpr = style.element.get_or_add_rPr()
        fonts = rpr.find(qn('w:rFonts'))
        if fonts is None:
            fonts = OxmlElement('w:rFonts')
            rpr.append(fonts)
        fonts.set(qn('w:eastAsia'), font_name)
        fonts.set(qn('w:ascii'), font_name)
        fonts.set(qn('w:hAnsi'), font_name)
    except Exception:
        pass

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # --- 封面页 ---
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(creator_name)
    _set_run_font(run, font_name, 36, color_rgb=(0x1A, 0x1A, 0x2E), bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("抖音视频文字稿合集")
    _set_run_font(run, font_name, 20, color_rgb=(0x66, 0x66, 0x66))

    doc.add_paragraph()

    total_videos = len(videos)
    transcribed = sum(1 for v in videos if v.get("transcript"))
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"共 {total_videos} 个视频 · 已转录 {transcribed} 个")
    _set_run_font(run, font_name, 12, color_rgb=(0x99, 0x99, 0x99))

    date_info = doc.add_paragraph()
    date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_info.add_run(f"生成日期：{datetime.now().strftime('%Y年%m月%d日')}")
    _set_run_font(run, font_name, 11, color_rgb=(0xAA, 0xAA, 0xAA))

    doc.add_page_break()

    # --- 目录页 ---
    toc_heading = doc.add_heading("目 录", level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(toc_heading, before_pt=0, after_pt=12)

    chapter_num = 0
    for video in videos:
        if video.get("transcript"):
            chapter_num += 1
            title_text = _clean_title(video.get("title", f"视频 {chapter_num}"))[:60]
            time_str = _format_time(video.get("create_time", ""))
            bookmark_name = f"chapter_{chapter_num}"

            toc_entry = doc.add_paragraph()
            _set_paragraph_spacing(toc_entry, before_pt=2, after_pt=2, line_spacing=1.8)

            # 添加可点击的超链接到对应章节
            link_text = f"{chapter_num}. {title_text}"
            if time_str:
                link_text += f"  （{time_str}）"
            _add_hyperlink_to_bookmark(toc_entry, bookmark_name, link_text, font_name, size_pt=10.5)

    doc.add_page_break()

    # --- 文字稿正文 ---
    if transcribed == 0:
        no_content = doc.add_paragraph()
        no_content.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = no_content.add_run("暂无已转录的文字稿内容")
        _set_run_font(run, font_name, 14, color_rgb=(0x99, 0x99, 0x99))

    chapter_num = 0
    for video in videos:
        transcript = video.get("transcript")
        if not transcript:
            continue

        chapter_num += 1
        title_text = _clean_title(video.get("title", f"视频 {chapter_num}"))[:80]
        bookmark_name = f"chapter_{chapter_num}"

        # 章节标题
        heading = doc.add_heading(f"{chapter_num}. {title_text}", level=2)
        _set_paragraph_spacing(heading, before_pt=18, after_pt=6)
        # 添加书签，目录链接指向这里
        _add_bookmark(heading, bookmark_name)

        # 元信息
        meta_parts = []
        time_str = _format_time(video.get("create_time", ""))
        if time_str:
            meta_parts.append(f"发布时间：{time_str}")
        if video.get("duration"):
            duration = video["duration"]
            if isinstance(duration, (int, float)) and duration > 0:
                minutes, seconds = divmod(int(duration), 60)
                meta_parts.append(f"时长：{minutes}:{seconds:02d}")

        if meta_parts:
            meta = doc.add_paragraph()
            _set_paragraph_spacing(meta, before_pt=0, after_pt=6, line_spacing=1.2)
            run = meta.add_run("  |  ".join(meta_parts))
            _set_run_font(run, font_name, 9, color_rgb=(0x99, 0x99, 0x99))

        # 文字稿内容 — 有时间戳分段则按段落显示，否则整段
        segments = transcript.get("segments", []) if isinstance(transcript, dict) else []
        text = transcript.get("text", "") if isinstance(transcript, dict) else str(transcript)

        if segments and len(segments) > 1:
            # 按时间戳分段显示，每段带时间标记
            for seg in segments:
                seg_para = doc.add_paragraph()
                _set_paragraph_spacing(seg_para, before_pt=1, after_pt=1, line_spacing=1.8)

                # 时间戳标记
                ts = _format_timestamp(seg.get("start", 0))
                ts_run = seg_para.add_run(f"[{ts}] ")
                _set_run_font(ts_run, font_name, 9, color_rgb=(0x99, 0x99, 0x99))

                # 文字内容
                seg_text = seg.get("text", "").strip()
                if seg_text:
                    text_run = seg_para.add_run(seg_text)
                    _set_run_font(text_run, font_name, 12)
        elif text:
            content_para = doc.add_paragraph()
            _set_paragraph_spacing(content_para, before_pt=6, after_pt=12, line_spacing=1.8)
            run = content_para.add_run(text)
            _set_run_font(run, font_name, 12)
            content_para.paragraph_format.first_line_indent = Pt(24)

        # 分隔线
        sep = doc.add_paragraph()
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(sep, before_pt=12, after_pt=12)
        run = sep.add_run("─" * 30)
        _set_run_font(run, font_name, 10, color_rgb=(0xDD, 0xDD, 0xDD))

    # 保存 — 文件名使用博主名称
    import re as _re
    safe_name = _re.sub(r'[\\/:*?"<>|]', '_', creator_name.strip() or sanitize_id(douyin_id))
    base_name = f"{safe_name}_文字稿合集"
    candidate = OUTPUT_DIR / f"{base_name}.docx"
    if candidate.exists():
        existing = list(OUTPUT_DIR.glob(f"{safe_name}_文字稿合集-*.docx"))
        max_num = 0
        for f in existing:
            match = _re.search(r'-(\d+)\.docx$', f.name)
            if match:
                max_num = max(max_num, int(match.group(1)))
        next_num = max_num + 1
        candidate = OUTPUT_DIR / f"{base_name}-{next_num}.docx"
    doc.save(str(candidate))
    return candidate


def _format_timestamp(seconds: float) -> str:
    """将秒数格式化为 MM:SS 时间戳"""
    seconds = max(0, int(seconds))
    minutes, secs = divmod(seconds, 60)
    if minutes >= 60:
        hours, minutes = divmod(minutes, 60)
        return f"{hours}:{minutes:02d}:{secs:02d}"
    return f"{minutes}:{secs:02d}"


def _format_time(time_val) -> str:
    """格式化时间显示"""
    if not time_val:
        return ""
    if isinstance(time_val, (int, float)):
        try:
            return datetime.fromtimestamp(time_val).strftime("%Y-%m-%d")
        except (OSError, ValueError, OverflowError):
            return ""
    if isinstance(time_val, str):
        return time_val[:10]
    return ""
